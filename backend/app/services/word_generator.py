"""
Geração de laudo de análise de risco em formato Word (.docx).
Espelha EXATAMENTE a estrutura do template laudo_analise_risco.html.
Seções, textos, tabelas e dados são idênticos ao PDF gerado.
"""
from __future__ import annotations
from datetime import datetime
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor

from app.engine.riscos import ComponentesRisco, RiscosConsolidados
from app.engine.avaliacao import ResultadoAvaliacao
from app.services.pdf_generator import EntradaLaudo, ProjetoInfo, ResponsavelTecnico

AZUL = RGBColor(0x1e, 0x3a, 0x8a)
CINZA_ESC = RGBColor(0x33, 0x41, 0x55)
CINZA = RGBColor(0x47, 0x55, 0x69)
CINZA_CLR = RGBColor(0x94, 0xa3, 0xb8)
VERMELHO = RGBColor(0xb9, 0x1c, 0x1c)
VERDE = RGBColor(0x15, 0x80, 0x3d)
BRANCO = RGBColor(0xff, 0xff, 0xff)
PRETO = RGBColor(0x0f, 0x17, 0x2a)
_AZUL_HEX = "1e3a8a"
_BRD_HEX = "e2e8f0"
_VERM_HEX = "fee2e2"
_VERDE_HEX = "dcfce7"

def _shd(cell, fill):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"),"clear")
    shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),fill); tcPr.append(shd)

def _brd(cell, color=_BRD_HEX, sz=4):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        el = OxmlElement(f"w:{side}"); el.set(qn("w:val"),"single")
        el.set(qn("w:sz"),str(sz)); el.set(qn("w:space"),"0"); el.set(qn("w:color"),color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def _p(doc, text="", bold=False, italic=False, size=10, color=None,
       align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text); run.bold=bold; run.italic=italic
        run.font.size=Pt(size); run.font.name="Calibri"
        if color: run.font.color.rgb=color

def _sec(doc, num, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(6)
    run = p.add_run(f"{num}. {title.upper()}")
    run.bold=True; run.font.size=Pt(11); run.font.color.rgb=AZUL; run.font.name="Calibri"
    pPr = p._p.get_or_add_pPr(); pBdr=OxmlElement("w:pBdr")
    bot=OxmlElement("w:bottom"); bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"8")
    bot.set(qn("w:space"),"1"); bot.set(qn("w:color"),_AZUL_HEX)
    pBdr.append(bot); pPr.append(pBdr)

def _sub(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(4)
    run=p.add_run(title); run.bold=True; run.font.size=Pt(10)
    run.font.color.rgb=AZUL; run.font.name="Calibri"

def _code(doc, lines):
    for line in lines:
        p = doc.add_paragraph(); p.paragraph_format.space_before=Pt(1)
        p.paragraph_format.space_after=Pt(1); p.paragraph_format.left_indent=Cm(0.5)
        run=p.add_run(line); run.font.name="Courier New"; run.font.size=Pt(9)
        run.font.color.rgb=CINZA_ESC
        pPr=p._p.get_or_add_pPr(); s=OxmlElement("w:shd")
        s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),"f1f5f9")
        pPr.append(s)

def _th(table, headers, bg=_AZUL_HEX):
    row = table.rows[0]
    for i,h in enumerate(headers):
        if i>=len(row.cells): break
        c=row.cells[i]; c.text=h; _shd(c,bg); _brd(c,"d1d5db")
        p=c.paragraphs[0]; p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT
        r=p.runs[0] if p.runs else p.add_run(h)
        r.bold=True; r.font.size=Pt(9); r.font.name="Calibri"
        r.font.color.rgb=BRANCO if bg==_AZUL_HEX else PRETO
        c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER

def _tr(table, idx, vals, bg=None, bold_first=False, mono_cols=None, right_cols=None, color_last=None):
    row=table.rows[idx]
    for i,v in enumerate(vals):
        if i>=len(row.cells): break
        c=row.cells[i]; c.text=v
        if bg: _shd(c,bg)
        _brd(c,"e5e7eb")
        p=c.paragraphs[0]
        if right_cols and i in right_cols: p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT
        r=p.runs[0] if p.runs else p.add_run(v)
        r.font.size=Pt(9); r.font.name="Calibri"
        if bold_first and i==0: r.bold=True
        if mono_cols and i in mono_cols: r.font.name="Courier New"
        if color_last and i==len(vals)-1: r.font.color.rgb=color_last
        c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER

def gerar_word_laudo(projeto, entrada, areas, eventos, componentes, riscos, avaliacao,
                     exige_protecao, responsavel, tem_spda=False, tem_dps=False,
                     medidas=None, linhas_info=None):
    doc = Document()
    sec = doc.sections[0]
    sec.page_height=Cm(29.7); sec.page_width=Cm(21.0)
    sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.0)
    sec.top_margin=Cm(2.0); sec.bottom_margin=Cm(2.0)
    doc.styles["Normal"].font.name="Calibri"; doc.styles["Normal"].font.size=Pt(10)

    hoje=datetime.now().strftime("%d/%m/%Y")
    codigo=f"PDA-{datetime.now().strftime('%Y%m%d%H%M%S')}"
    detalhes = riscos.detalhes or {}
    R1 = riscos.R1
    F = detalhes.get("F", detalhes.get("F_global", 0.0)) if detalhes else 0.0
    R3 = getattr(riscos, "R3", 0.0) or detalhes.get("R3", detalhes.get("R3_global", 0.0))
    R4 = getattr(riscos, "R4", 0.0) or detalhes.get("R4", detalhes.get("R4_global", 0.0))
    tem_R3 = bool(detalhes.get("tem_R3")) or R3 > 0
    tem_R4 = bool(detalhes.get("tem_R4")) or R4 > 0
    FT = float(detalhes.get("FT_global", detalhes.get("FT", getattr(riscos, "FT", 0.1))) or 0.1)
    f_ok = bool(detalhes.get("F_atende", detalhes.get("F_conforme", F <= FT)))
    r1_ok=R1<=1e-5; r3_ok=R3<=1e-4; r4_ok=R4<=1e-3
    painel_resultados=[
        ("R1 — Proteção à Vida", R1, "1,00 x 10^-5", r1_ok, "R1", False),
        ("F — Frequência de danos", F, f"{FT:.2e}", f_ok, "F", False),
    ]
    if tem_R3:
        painel_resultados.append(("R3 — Patrimônio Cultural", R3, "1,00 x 10^-4", r3_ok, "R3", False))
    if tem_R4:
        painel_resultados.append(("R4 — Perdas Econômicas (informativo)", R4, "1,00 x 10^-3", r4_ok, "R4", True))
    geral_aprovado = all(item[3] for item in painel_resultados if not item[5])
    falhas_conformidade = [item for item in painel_resultados if (not item[3]) and (not item[5])]
    AD=areas.get("AD",0); AM=areas.get("AM",0); AL=areas.get("AL",0); AI=areas.get("AI",0)
    ND=eventos.get("ND",0); NM=eventos.get("NM",0); NL=eventos.get("NL",0); NI=eventos.get("NI",0)
    medidas=medidas or {}
    rs_valor = float(getattr(entrada, "rS", 1.0) or 1.0)
    rs_label = getattr(entrada, "tipo_construcao_label", "Robusta / alvenaria-concreto") or "Robusta / alvenaria-concreto"

    # CAPA
    p_topo=doc.add_paragraph()
    p_topo.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p_topo.paragraph_format.space_before=Pt(0); p_topo.paragraph_format.space_after=Pt(20)
    pPr=p_topo._p.get_or_add_pPr(); pBdr=OxmlElement("w:pBdr")
    top_el=OxmlElement("w:top"); top_el.set(qn("w:val"),"single"); top_el.set(qn("w:sz"),"48")
    top_el.set(qn("w:space"),"1"); top_el.set(qn("w:color"),_AZUL_HEX)
    pBdr.append(top_el); pPr.append(pBdr)
    r=p_topo.add_run("C O N F O R M E  À  A B N T  N B R  5 4 1 9 - 2 : 2 0 2 6")
    r.font.size=Pt(7.5); r.bold=True; r.font.color.rgb=AZUL; r.font.name="Calibri"

    p2=doc.add_paragraph(); p2.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER; p2.paragraph_format.space_after=Pt(4)
    r2=p2.add_run("MEMORIAL DE CÁLCULO"); r2.font.size=Pt(12); r2.bold=True; r2.font.color.rgb=AZUL; r2.font.name="Calibri"
    p3=doc.add_paragraph(); p3.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER; p3.paragraph_format.space_before=Pt(4); p3.paragraph_format.space_after=Pt(4)
    r3=p3.add_run("Análise de Risco"); r3.font.size=Pt(28); r3.bold=True; r3.font.color.rgb=AZUL; r3.font.name="Calibri"
    p4=doc.add_paragraph(); p4.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER; p4.paragraph_format.space_after=Pt(20)
    r4=p4.add_run("SISTEMA E REDE DE PROTEÇÃO CONTRA DESCARGAS ATMOSFÉRICAS (SPDA)"); r4.font.size=Pt(9); r4.font.color.rgb=CINZA; r4.font.name="Calibri"

    capa_rows=[("IDENTIFICAÇÃO DA OBRA",projeto.nome or "—"),("AUTORIA TÉCNICA",responsavel.nome or "—"),
               ("REGISTRO PROFISSIONAL",responsavel.registro or "Conselho nº ________"),
               ("ART / RRT / TRT EMISSÃO",responsavel.art or "________"),
               ("DATA DE EXPEDIÇÃO",hoje),("CÓDIGO DE RASTREIO",codigo)]
    tbl_capa=doc.add_table(rows=len(capa_rows),cols=2); tbl_capa.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,(label,val) in enumerate(capa_rows):
        row=tbl_capa.rows[i]; c0,c1=row.cells[0],row.cells[1]; c0.text=label; c1.text=val
        _shd(c0,"dbeafe"); _brd(c0,_BRD_HEX); _brd(c1,_BRD_HEX)
        p0=c0.paragraphs[0]; p0.runs[0].bold=True; p0.runs[0].font.size=Pt(9); p0.runs[0].font.color.rgb=AZUL; p0.runs[0].font.name="Calibri"
        p1=c1.paragraphs[0]
        if p1.runs: p1.runs[0].font.size=Pt(10); p1.runs[0].font.name="Calibri"
    doc.add_paragraph()
    p_rod=doc.add_paragraph(); p_rod.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r_rod=p_rod.add_run("Documento elaborado estritamente conforme as diretrizes da ABNT NBR 5419-2:2026 — Análise de Risco.")
    r_rod.font.size=Pt(8.5); r_rod.italic=True; r_rod.font.color.rgb=CINZA_CLR; r_rod.font.name="Calibri"
    doc.add_page_break()

    # ÍNDICE
    p_idx=doc.add_paragraph(); p_idx.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER; p_idx.paragraph_format.space_after=Pt(12)
    r_idx=p_idx.add_run("Índice Analítico"); r_idx.bold=True; r_idx.font.size=Pt(14); r_idx.font.color.rgb=AZUL; r_idx.font.name="Calibri"
    secs=[("1","Propósito e Delimitação"),("2","Base Normativa Aplicada"),("3","Parâmetros Construtivos da Edificação"),
          ("3.1","Coordenadas e Dimensões"),("3.2","Modelagem e Premissas de Cálculo"),("4","Variáveis Atmosféricas e Locais"),
          ("4.1","Georreferenciamento"),("4.2","Índice Ceráunico (NG)"),("4.3","Projeções de Incidência"),
          ("5","Matriz de Soluções Mitigadoras"),("6","Condutores e Redes Importadas"),("7","Edificações em Anexo"),
          ("8","Superfícies de Interceptação"),("9","Estimativa de Sinistros Anuais"),("10","Mapeamento de Zonas de Segurança"),
          ("11","Dicionário de Variáveis de Dano"),("12","Fórmula de Consolidação dos Limites"),
          ("13","Painel de Resultados Gerais"),("14","Diagnóstico Detalhado — Risco à Vida (R1)"),
          ("15","Diagnóstico Detalhado — Falhas Sistêmicas (F)"),("16","Resumo Computacional — Setor 01"),("17","Parecer Técnico Conclusivo")]
    for num,titulo in secs:
        p_s=doc.add_paragraph(); p_s.paragraph_format.space_before=Pt(1); p_s.paragraph_format.space_after=Pt(1)
        indent="." in num; p_s.paragraph_format.left_indent=Cm(1.0 if indent else 0)
        r_s=p_s.add_run(f"{num}. {titulo}"); r_s.font.size=Pt(9 if not indent else 8.5)
        r_s.font.color.rgb=CINZA_ESC if not indent else CINZA; r_s.font.name="Calibri"
    doc.add_page_break()

    # SEÇÃO 1
    _sec(doc,"1","Propósito e Delimitação")
    _p(doc,"O escopo deste laudo consiste em fundamentar matematicamente a Análise de Risco associada à infraestrutura, utilizando estritamente os vetores de cálculo definidos pela atualizada ABNT NBR 5419-2:2026 (Proteção contra descargas atmosféricas — Parte 2: Análise de risco).",size=9.5)
    _p(doc,"O texto normativo dita a metodologia para quantificar as prováveis perdas ocasionadas por raios (sejam impactos diretos ou irradiações magnéticas) sobre a edificação. A partir deste diagnóstico numérico, o estudo orienta a adoção de medidas mitigadoras para manter os níveis de segurança dentro da zona de tolerabilidade aceitável.",size=9.5)
    _p(doc,"A esteira de raciocínio adotada na avaliação engloba as seguintes etapas operacionais:",size=9.5)
    for item in ["Divisão e isolamento das zonas protegidas (ZPR) do edifício;","Mapeamento quantitativo das variáveis de risco para cada setor;","Soma matricial das probabilidades de perdas patrimoniais e humanas;","Confronto direto dos resultados contra os tetos de corte legais (RT e FT);","Dimensionamento do arcabouço de proteção necessário."]:
        p_li=doc.add_paragraph(); p_li.paragraph_format.left_indent=Cm(1); p_li.paragraph_format.space_after=Pt(2)
        r_li=p_li.add_run(f"  •  {item}"); r_li.font.size=Pt(9.5); r_li.font.color.rgb=CINZA_ESC; r_li.font.name="Calibri"
    _sub(doc,"1.1 Dicionário de Constantes e Matrizes")
    _code(doc,["NG = Densidade de raios no solo (descargas/km2/ano) mapeada [Tab. F.1]","AD = Projecao de cobertura direta da estrutura (m2) [Eq. A.1]","AM = Perimetro de influencia magnetica (m2)","CD = Multiplicador de vizinhanca topografica (adim.) [Tab. A.1]","ND = Expectativa anual de incidencias diretas (1/ano) [Eq. A.3]","NM = Expectativa anual de descargas orbitais (1/ano) [Eq. A.5]","NL = Choques transferidos pelas linhas (1/ano) [Eq. A.7]","NI = Inducoes por quedas proximas as linhas (1/ano) [Eq. A.9]","PX = Probabilidade de dano associada ao evento perigoso [Anexo B]","rS = Fator associado ao tipo de construcao [Tab. C.7]","RT = Risco toleravel para R1/R3 [Tab. 4]","FT = Teto de frequencia de paralisacoes [Secao 7 - Atualizacao 2026]"])

    # SEÇÃO 2
    _sec(doc,"2","Base Normativa Aplicada")
    _p(doc,"Toda a lógica computacional e os arranjos técnicos sugeridos baseiam-se na seguinte bibliografia padronizada:",size=9.5)
    for ref in ["ABNT NBR 5419-1:2026 — Princípios gerais físicos e conceituais do sistema.","ABNT NBR 5419-2:2026 — Gerenciamento e cálculo estrutural de riscos.","ABNT NBR 5419-3:2026 — Salvaguarda contra colapso físico e fatalidades em humanos.","ABNT NBR 5419-4:2026 — Blindagem de malhas e equipamentos eletroeletrônicos sensíveis.","Série IEC 62305 (2024) — Consenso e diretriz técnica internacional para fenômenos atmosféricos.","ABNT NBR 5410:2004 — Práticas seguras em instalações elétricas de baixa tensão."]:
        p_ref=doc.add_paragraph(); p_ref.paragraph_format.left_indent=Cm(0.5); p_ref.paragraph_format.space_after=Pt(3)
        r_ref=p_ref.add_run(f"▪  {ref}"); r_ref.font.size=Pt(9.5); r_ref.font.color.rgb=CINZA; r_ref.font.name="Calibri"

    # SEÇÃO 3
    doc.add_page_break()
    _sec(doc,"3","Parâmetros Construtivos da Edificação")
    _sub(doc,"3.1 Coordenadas e Dimensões")
    _p(doc,"Os atributos arquitetônicos e físicos que definem o contorno da obra em estudo estão consolidados no quadro abaixo:",size=9.5)
    tbl3=doc.add_table(rows=10,cols=4); tbl3.style="Table Grid"
    _th(tbl3,["Atributo","Significado Técnico","Medição Final","Origem do Dado"])
    rows3=[("L","Eixo Longitudinal (m)",f"{entrada.L:.1f} m","Planta da Edificação"),("W","Eixo Transversal (m)",f"{entrada.W:.1f} m","Planta da Edificação"),("H","Cota Máxima de Telhado (m)",f"{entrada.H:.1f} m","Planta da Edificação"),("Hp","Elementos de cumeada e caixas d'água (m)","—","Inspeção Visual"),("CD","Impacto do entorno (relevo)",f"{entrada.CD:.2f}","Tabela A.1"),("PB","Barreira de interceptação SPDA",medidas.get("spda_nivel","NENHUM"),"Tabela B.2"),("rS","Fator associado ao tipo de construção",f"{rs_valor:.3g} — {rs_label}","Tabela C.7"),("LF","Potencial de desabamento/fogo",entrada.tipo_estrutura_label or str(entrada.tipo_estrutura),"Tabela C.2"),("NG","Incidência geolocalizada",f"{entrada.NG:.1f} raios/km²/ano","Tabela F.1")]
    for i,r in enumerate(rows3,1): _tr(tbl3,i,list(r),bold_first=True)
    _sub(doc,"3.2 Modelagem e Premissas de Cálculo")
    _p(doc,"Para o dimensionamento correto das áreas de captura (AD e AM), a edificação foi abstraída geometricamente como um bloco sólido de teto plano. O raio de interceptação acompanha a sombra tridimensional gerada pela altura total do prédio.",size=9.5)
    _code(doc,[f"[Eq. A.1] Sombra de Captura Direta:",f"AD = L x W + 2 x (3H) x (L + W) + pi x (3H)^2",f"AD = {entrada.L:.1f} x {entrada.W:.1f} + 2 x (3x{entrada.H:.1f}) x ({entrada.L:.1f} + {entrada.W:.1f}) + pi x (3x{entrada.H:.1f})^2",f"AD = {AD:.3f} m^2","",f"[Eq. A.6] Área de exposição para descargas próximas à estrutura:",f"AM = 2 x 500 x (L + W) + pi x 500^2",f"AM = {AM:.3f} m^2"])
    _p(doc,f"O diagnóstico topográfico determina o coeficiente {entrada.localizacao_label or str(entrada.localizacao)}.",size=9.5)

    # SEÇÃO 4
    _sec(doc,"4","Variáveis Atmosféricas e Locais")
    _sub(doc,"4.1 Georreferenciamento")
    _p(doc,f"Considerando a implantação geográfica do projeto em {getattr(projeto,'endereco','') or 'Não preenchido'}, o avaliador insere o fator ambiental de CD = {entrada.CD:.2f} (caracterizando a vizinhança geomorfológica: {entrada.localizacao_label or str(entrada.localizacao)}).",size=9.5)
    _sub(doc,"4.2 Índice Ceráunico (NG)")
    _p(doc,"A taxa de incidência de raios tocando o solo na região obedece à mais recente tabulação do painel meteorológico (Tabela F.1 da NBR 5419-2:2026). O número base imposto aos cálculos estatísticos é:",size=9.5)
    p_ng=doc.add_paragraph(); p_ng.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER; p_ng.paragraph_format.space_before=Pt(6); p_ng.paragraph_format.space_after=Pt(2)
    r_ng=p_ng.add_run(f"{entrada.NG:.1f}"); r_ng.font.size=Pt(28); r_ng.bold=True; r_ng.font.color.rgb=AZUL; r_ng.font.name="Calibri"
    p_ng2=doc.add_paragraph(); p_ng2.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER; p_ng2.paragraph_format.space_after=Pt(8)
    r_ng2=p_ng2.add_run("DENSIDADE MÉDIA ACUMULADA (RAIOS / KM² / ANO)"); r_ng2.font.size=Pt(8); r_ng2.font.color.rgb=CINZA_CLR; r_ng2.font.name="Calibri"
    _sub(doc,"4.3 Projeções de Incidência")
    _code(doc,[f"[Eq. A.3] Raios Acertando a Cobertura:","ND = NG x AD x CD x 10^-6",f"ND = {entrada.NG:.2f} x {AD:.3f} x {entrada.CD:.2f} x 10^-6 = {ND:.3e} interacoes/ano","",f"[Eq. A.5] Raios Caindo no Terreno Adjacente:","NM = NG x AM x 10^-6",f"NM = {entrada.NG:.2f} x {AM:.3f} x 10^-6 = {NM:.3e} interacoes/ano"])

    # SEÇÃO 5
    _sec(doc,"5","Matriz de Soluções Mitigadoras")
    _p(doc,"Os recursos mecânicos e elétricos de segurança instalados (ou previstos) atuam diretamente na atenuação das probabilidades de dano, freando a letalidade do raio. O layout abaixo retrata as barreiras avaliadas:",size=9.5)
    tbl5=doc.add_table(rows=7,cols=3); tbl5.style="Table Grid"
    _th(tbl5,["Equipamento / Estratégia","Embasamento Técnico","Configuração Aplicada"])
    mitig5=[("Arranjo de Captores (PB)","Rede externa de blindagem (Tab. B.2)",medidas.get("spda_nivel","NENHUM")),("DPS Distribuídos (PSPD)","Cascatas internas de proteção (Tab. B.3)",medidas.get("dps_coordenados_nivel","NENHUM")),("DPS Principal (PEB)","Fronteira do quadro central (Tab. B.7)",medidas.get("dps_classe_I_entrada","NENHUM")),("Demarcações Visuais (PTA)","Inibição de toque/passo (Tab. B.1)","Incluído" if medidas.get("aviso_alerta") else "Não Aplicado"),("Afastamento Dielétrico","Descolamento das descidas da parede (§B.3.2)","Em conformidade" if medidas.get("isolacao_eletrica") else "Padrão Ausente"),("Grade de Aterramento (TAP)","Ligação equipotencial unificada","Garantido" if medidas.get("malha_equipotencializacao") else "Inexistente")]
    for i,m in enumerate(mitig5,1): _tr(tbl5,i,list(m))

    # SEÇÃO 6
    _sec(doc,"6","Condutores e Redes Importadas")
    _p(doc,"Os dutos de energia, lógica e fibra que furam as barreiras do edifício atuam como antenas, transportando surtos de vizinhança (S3 e S4).",size=9.5)
    _code(doc,["Modelagem das Linhas de Utilidade:","[Eq. A.7] NL = NG x AL x CI x CT x CE x 10^-6","[Eq. A.9] NI = NG x AI x CI x CT x CE x 10^-6","Onde os poligonos sao: AL = 40 x LL | AI = 4000 x LL"])
    if NL>0 or NI>0:
        _sub(doc,"6.1 Entradas de Concessionária / Alimentadores")
        tbl6=doc.add_table(rows=3,cols=2); tbl6.style="Table Grid"
        _th(tbl6,["Sigla de Evento","Resultado Estatístico"])
        _tr(tbl6,1,["NL (Raio caindo diretamente nos fios da rua)",f"{NL:.3e} disturbios/ano"],mono_cols=[1],right_cols=[1])
        _tr(tbl6,2,["NI (Pulso magnetico por raio perto do poste)",f"{NI:.3e} disturbios/ano"],mono_cols=[1],right_cols=[1])

    # SEÇÃO 8
    doc.add_page_break()
    _sec(doc,"8","Superfícies de Interceptação")
    rows8=[("Prédio Avaliado","AD (Alvo Direto)",f"{AD:.1f}"),("Região no Entorno","AM (Coroa Magnética)",f"{AM:.1f}")]
    if AL>0: rows8.extend([("Cabeamento Externo","AL (Rede Captadora)",f"{AL:.1f}"),("Zona Vizinha aos Cabos","AI (Zona de Indução)",f"{AI:.1f}")])
    tbl8=doc.add_table(rows=len(rows8)+1,cols=3); tbl8.style="Table Grid"
    _th(tbl8,["Ponto de Foco","Variável de Abrangência","Dimensão Computada (m²)"])
    for i,r in enumerate(rows8,1): _tr(tbl8,i,list(r),right_cols=[2],mono_cols=[2])

    # SEÇÃO 12
    _sec(doc,"12","Fórmula de Consolidação dos Limites")
    _p(doc,"A arquitetura do cálculo junta todas as variáveis probabilísticas de falha em dois indicadores decisivos:",size=9.5)
    _code(doc,["Potencial de Morte e Lesões (R1):","R1 = RA + RB + RU + RV + RC* + RM* + RW* + RZ*","* Componentes D3 aplicáveis somente quando falhas de sistemas internos puderem colocar imediatamente em risco a vida humana ou o meio ambiente.","","Métrica de Dano e Paralisação (F) [Revisão 2026]:","F = FB + FC + FM + FV + FW + FZ"])
    _sub(doc,"Limites de aceitação aplicáveis (Tab. 4 e Seção 7)")
    tbl12=doc.add_table(rows=3,cols=3); tbl12.style="Table Grid"
    _th(tbl12,["Métrica","Natureza da Consequência","Barreira Numérica"])
    _tr(tbl12,1,["R1","Sinistros com Vítimas Fatais ou Graves","1,00 x 10^-5"],mono_cols=[2],right_cols=[2],color_last=VERMELHO)
    _tr(tbl12,2,["F","Frequência de danos dos sistemas internos",f"{FT:.2e}"],mono_cols=[2],right_cols=[2],color_last=VERMELHO)

    # SEÇÃO 13
    doc.add_page_break()
    _sec(doc,"13","Painel de Resultados Gerais")
    _p(doc,"Abaixo, apresentamos os consolidadores analíticos que demonstram o desempenho da edificação perante as exigências normativas vigentes:",size=9.5)
    tbl13=doc.add_table(rows=3,cols=4); tbl13.style="Table Grid"
    _th(tbl13,["Indicador","Valor Calculado","Limite Legal","Status"])
    for i,(label,val,rt,ok) in enumerate([(f"Integridade Humana (R1)",R1,"1,00 x 10^-5",r1_ok),(f"Índice de Falhas (F)",F,f"{FT:.2e}",f_ok)],1):
        row=tbl13.rows[i]
        vals=[label,f"{val:.2e}",rt,"APROVADO" if ok else "INSUFICIENTE"]
        for j,v in enumerate(vals):
            c=row.cells[j]; c.text=v
            if j==3: _shd(c,_VERDE_HEX if ok else _VERM_HEX)
            _brd(c,"e5e7eb"); p=c.paragraphs[0]
            r=p.runs[0] if p.runs else p.add_run(v)
            r.font.size=Pt(9); r.font.name="Calibri"
            if j==1: r.font.name="Courier New"; r.bold=True
            if j==3: r.bold=True; r.font.color.rgb=VERDE if ok else VERMELHO
    doc.add_paragraph()
    for label_v,val_v,ok_v,rt_label in [("AMEAÇA À VIDA (R1)",R1,r1_ok,"1,00 × 10⁻⁵"),("TAXA DE FALHAS (F)",F,f_ok,f"{FT:.2e}")]:
        p_s=doc.add_paragraph(); p_s.paragraph_format.space_before=Pt(8); p_s.paragraph_format.space_after=Pt(2)
        rs1=p_s.add_run(f"{label_v}   "); rs1.bold=True; rs1.font.size=Pt(10); rs1.font.color.rgb=AZUL; rs1.font.name="Calibri"
        rs2=p_s.add_run(f"{val_v:.2e}"); rs2.bold=True; rs2.font.size=Pt(14); rs2.font.color.rgb=VERDE if ok_v else VERMELHO; rs2.font.name="Courier New"
        p_s2=doc.add_paragraph(); p_s2.paragraph_format.space_after=Pt(2)
        rs3=p_s2.add_run("  APROVADO" if ok_v else "  INSUFICIENTE"); rs3.bold=True; rs3.font.size=Pt(11); rs3.font.color.rgb=VERDE if ok_v else VERMELHO; rs3.font.name="Calibri"
        p_s3=doc.add_paragraph(); p_s3.paragraph_format.space_after=Pt(8)
        rs4=p_s3.add_run(f"  Bloqueio Normativo: {rt_label}"); rs4.font.size=Pt(9); rs4.font.color.rgb=CINZA_CLR; rs4.font.name="Calibri"

    # SEÇÃO 14
    doc.add_page_break()
    _sec(doc,"14","Diagnóstico Detalhado — Risco à Vida (R1)")
    r1_comps_d=[("Explosão Estrutural (RB)",componentes.RB),("Fogo Transmitido (RV)",componentes.RV),("Choque por DDP (RA)",componentes.RA),("Choque na Fiação (RU)",componentes.RU),("Pane LEMP Central (RC)",componentes.RC),("Onda LEMP Externa (RM)",componentes.RM),("Surto Conduzido (RW)",componentes.RW),("Campo Induzido (RZ)",componentes.RZ)]
    comps_pos=[(l,v) for l,v in r1_comps_d if v>0]
    _sub(doc,"Radiografia dos Agravantes de R1")
    if comps_pos:
        tbl14a=doc.add_table(rows=len(comps_pos)+1,cols=2); tbl14a.style="Table Grid"
        _th(tbl14a,["Componente","Valor"])
        for i,(label,val) in enumerate(comps_pos,1): _tr(tbl14a,i,[label,f"{val:.2e}"],mono_cols=[1],right_cols=[1])
    _sub(doc,"Espelho de Submúltiplos (R1)")
    comp_rows=[("[S1] Raio no Prédio","RA (Descarga pelo chão ou paredes)",componentes.RA),("[S1] Raio no Prédio","RB (Fagulhas que geram chamas)",componentes.RB),("[S1] Raio no Prédio","RC (Picos que queimam aparelhos vitais)",componentes.RC),("[S2] Raio no Vizinho","RM (Irradiação que interfere em CPUs)",componentes.RM),("[S3] Raio na Rede Externa","RU (Choque que viaja pela tomada)",componentes.RU),("[S3] Raio na Rede Externa","RV (Explosão de quadros elétricos)",componentes.RV),("[S3] Raio na Rede Externa","RW (Surto severo que torra placas)",componentes.RW),("[S4] Raio perto da Rede","RZ (Interferência induzida sutil)",componentes.RZ)]
    tbl14b=doc.add_table(rows=len(comp_rows)+1,cols=3); tbl14b.style="Table Grid"
    _th(tbl14b,["Gatilho do Evento","Manifestação do Dano","Peso Calculado"])
    for i,(g,m,v) in enumerate(comp_rows,1): _tr(tbl14b,i,[g,m,f"{v:.3e}"],mono_cols=[2],right_cols=[2])
    doc.add_paragraph()
    p_tot=doc.add_paragraph()
    r_tot=p_tot.add_run(f"ÍNDICE FINAL APURADO (R1):    {R1:.3e}")
    r_tot.bold=True; r_tot.font.size=Pt(11); r_tot.font.color.rgb=VERMELHO if not r1_ok else VERDE; r_tot.font.name="Courier New"

    # SEÇÃO 15 — Frequência completa
    doc.add_page_break()
    _sec(doc,"15","Frequência de Danos — Memória Completa")
    _p(doc,"Componentes calculados pela API principal: F = FB + FC + FM + FV + FW + FZ.",size=9.5)
    detalhes = riscos.detalhes or {}
    zonas_calc = detalhes.get("zonas", []) or []
    linhas_calc = detalhes.get("linhas", []) or []

    if zonas_calc:
        _sub(doc,"Componentes por zona")
        tbl15=doc.add_table(rows=len(zonas_calc)+1,cols=8); tbl15.style="Table Grid"
        _th(tbl15,["Zona","FB","FC","FM","FV","FW","FZ","F total"])
        for i,z in enumerate(zonas_calc,1):
            _tr(tbl15,i,[z.get("nome","—"),f"{z.get('FB',0):.3e}",f"{z.get('FC',0):.3e}",f"{z.get('FM',0):.3e}",f"{z.get('FV',0):.3e}",f"{z.get('FW',0):.3e}",f"{z.get('FZ',0):.3e}",f"{z.get('F',0):.3e}"],mono_cols=[1,2,3,4,5,6,7],right_cols=[1,2,3,4,5,6,7])

        _sub(doc,"Parâmetros PM/PMS por zona")
        tbl15b=doc.add_table(rows=len(zonas_calc)+1,cols=7); tbl15b.style="Table Grid"
        _th(tbl15b,["Zona","PC","PMS","PM","KS1","KS2","KS4"])
        for i,z in enumerate(zonas_calc,1):
            _tr(tbl15b,i,[z.get("nome","—"),f"{z.get('PC_calc',0):.3e}",f"{z.get('PMS_calc',0):.3e}",f"{z.get('PM_calc',0):.3e}",f"{z.get('KS1_calc',0):.3e}",f"{z.get('KS2_calc',0):.3e}",f"{z.get('KS4_calc',0):.3e}"],mono_cols=[1,2,3,4,5,6],right_cols=[1,2,3,4,5,6])

        _sub(doc,"Contribuição das linhas em FV/FW/FZ")
        for z in zonas_calc:
            _p(doc,z.get("nome","Zona"),bold=True,size=9.5,color=AZUL)
            linhas_contrib=z.get("linhas_contrib",[]) or []
            tbl15c=doc.add_table(rows=len(linhas_contrib)+1,cols=4); tbl15c.style="Table Grid"
            _th(tbl15c,["Linha","FV","FW","FZ"])
            for i,lc in enumerate(linhas_contrib,1):
                _tr(tbl15c,i,[lc.get("nome","—"),f"{lc.get('FV',0):.3e}",f"{lc.get('FW',0):.3e}",f"{lc.get('FZ',0):.3e}"],mono_cols=[1,2,3],right_cols=[1,2,3])

    # SEÇÃO 16 — Patrimônio cultural
    doc.add_page_break()
    _sec(doc,"16","Perda de Patrimônio Cultural — R3")
    _p(doc,"Quando houver patrimônio cultural, são apresentados RB3, RV3 e R3 por zona.",size=9.5)
    r3_zonas=[z for z in zonas_calc if (z.get("R3",0) or z.get("RB3",0) or z.get("RV3",0))]
    tbl16=doc.add_table(rows=max(len(r3_zonas),1)+1,cols=4); tbl16.style="Table Grid"
    _th(tbl16,["Zona","RB3","RV3","R3"])
    if r3_zonas:
        for i,z in enumerate(r3_zonas,1):
            _tr(tbl16,i,[z.get("nome","—"),f"{z.get('RB3',0):.3e}",f"{z.get('RV3',0):.3e}",f"{z.get('R3',0):.3e}"],mono_cols=[1,2,3],right_cols=[1,2,3])
    else:
        _tr(tbl16,1,["Não informado / não aplicável","0.000e+00","0.000e+00","0.000e+00"],mono_cols=[1,2,3],right_cols=[1,2,3])

    # SEÇÃO 17 — Perdas econômicas
    doc.add_page_break()
    _sec(doc,"17","Perdas Econômicas — R4")
    _p(doc,"Quando habilitada a avaliação econômica, são apresentados todos os componentes R4 por zona e por linha. O R4 é informativo/opcional no Anexo D e não reprova o parecer normativo principal.",size=9.5)
    r4_zonas=[z for z in zonas_calc if (z.get("R4",0) or z.get("RA4",0) or z.get("RB4",0) or z.get("RC4",0) or z.get("RM4",0) or z.get("RU4",0) or z.get("RV4",0) or z.get("RW4",0) or z.get("RZ4",0))]
    tbl17=doc.add_table(rows=max(len(r4_zonas),1)+1,cols=10); tbl17.style="Table Grid"
    _th(tbl17,["Zona","RA4","RB4","RC4","RM4","RU4","RV4","RW4","RZ4","R4"])
    if r4_zonas:
        for i,z in enumerate(r4_zonas,1):
            _tr(tbl17,i,[z.get("nome","—"),f"{z.get('RA4',0):.3e}",f"{z.get('RB4',0):.3e}",f"{z.get('RC4',0):.3e}",f"{z.get('RM4',0):.3e}",f"{z.get('RU4',0):.3e}",f"{z.get('RV4',0):.3e}",f"{z.get('RW4',0):.3e}",f"{z.get('RZ4',0):.3e}",f"{z.get('R4',0):.3e}"],mono_cols=[1,2,3,4,5,6,7,8,9],right_cols=[1,2,3,4,5,6,7,8,9])
    else:
        _tr(tbl17,1,["Não informado / não aplicável"]+["0.000e+00"]*9,mono_cols=list(range(1,10)),right_cols=list(range(1,10)))

    # SEÇÃO 18 — Linhas externas
    doc.add_page_break()
    _sec(doc,"18","Linhas Externas — Eventos Calculados")
    tbl18=doc.add_table(rows=max(len(linhas_calc),1)+1,cols=6); tbl18.style="Table Grid"
    _th(tbl18,["Linha","AL","AI","NL","NI","NDJ"])
    if linhas_calc:
        for i,l in enumerate(linhas_calc,1):
            nome=(l.get("id","") + " — " + l.get("nome","—")).strip(" —")
            _tr(tbl18,i,[nome,f"{l.get('AL_total',0):.1f}",f"{l.get('AI_total',0):.1f}",f"{l.get('NL_total',0):.3e}",f"{l.get('NI_total',0):.3e}",f"{l.get('NDJ',0):.3e}"],mono_cols=[1,2,3,4,5],right_cols=[1,2,3,4,5])
    else:
        _tr(tbl18,1,["Sem linhas externas","0.0","0.0","0.000e+00","0.000e+00","0.000e+00"],mono_cols=[1,2,3,4,5],right_cols=[1,2,3,4,5])


    # PARECER
    doc.add_page_break()
    _sec(doc,"19","Parecer Técnico Conclusivo")
    _p(doc,"A partir do escrutínio numérico e da comparação dos resultados contra a tábua de tolerância da ABNT NBR 5419-2:2026, emite-se a seguinte posição técnica acerca da conformidade desta obra:",size=9.5)
    tbl_fin=doc.add_table(rows=len(painel_resultados)+1,cols=4); tbl_fin.style="Table Grid"
    _th(tbl_fin,["Pilar de Análise","Métrica Levantada","Limite / Referência","Status Final"])
    for i,(label,val,rt,ok,cod,informativo) in enumerate(painel_resultados,1):
        row=tbl_fin.rows[i]
        status = "INFORMATIVO" if informativo else ("HABILITADO" if ok else "FORA DE NORMA")
        vals=[label,f"{val:.3e}",rt,status]
        for j,v in enumerate(vals):
            c=row.cells[j]; c.text=v
            if j==3: _shd(c,_VERDE_HEX if (ok or informativo) else _VERM_HEX)
            _brd(c,"e5e7eb"); p=c.paragraphs[0]; r=p.runs[0] if p.runs else p.add_run(v)
            r.font.size=Pt(9); r.font.name="Calibri"
            if j==1: r.font.name="Courier New"; r.bold=True
            if j==3: r.bold=True; r.font.color.rgb=VERDE if (ok or informativo) else VERMELHO
    doc.add_paragraph()
    p_conc=doc.add_paragraph(); p_conc.paragraph_format.space_before=Pt(8)
    pPr_c=p_conc._p.get_or_add_pPr(); shd_c=OxmlElement("w:shd")
    shd_c.set(qn("w:val"),"clear"); shd_c.set(qn("w:color"),"auto"); shd_c.set(qn("w:fill"),"f0fdf4" if geral_aprovado else "fef2f2"); pPr_c.append(shd_c)
    if geral_aprovado:
        r_conc=p_conc.add_run("APROVAÇÃO EMITIDA: TODOS OS INDICADORES APLICÁVEIS ATENDEM AOS LIMITES TOLERÁVEIS.\n"); r_conc.bold=True; r_conc.font.size=Pt(11); r_conc.font.color.rgb=VERDE; r_conc.font.name="Calibri"
        r_conc2=p_conc.add_run("O crivo matemático evidencia que os indicadores avaliados permanecem abaixo dos limites toleráveis aplicáveis definidos pela ABNT NBR 5419-2:2026."); r_conc2.font.size=Pt(9.5); r_conc2.font.color.rgb=CINZA_ESC; r_conc2.font.name="Calibri"
    else:
        r_conc=p_conc.add_run("APROVAÇÃO NÃO EMITIDA: EXISTEM PARÂMETROS FORA DOS LIMITES TOLERÁVEIS.\n"); r_conc.bold=True; r_conc.font.size=Pt(11); r_conc.font.color.rgb=VERMELHO; r_conc.font.name="Calibri"
        falhas = ", ".join([f[4] for f in falhas_conformidade]) or "indicadores críticos"
        r_conc2=p_conc.add_run(f"A análise identificou parâmetro(s) fora de norma: {falhas}. Devem ser adotadas medidas adicionais e a análise deve ser recalculada após a revisão das condições de proteção."); r_conc2.font.size=Pt(9.5); r_conc2.font.color.rgb=CINZA_ESC; r_conc2.font.name="Calibri"

    # ASSINATURA
    doc.add_paragraph()
    p_sig=doc.add_paragraph(); p_sig.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER; p_sig.paragraph_format.space_before=Pt(24)
    pPr_sig=p_sig._p.get_or_add_pPr(); pBdr_sig=OxmlElement("w:pBdr")
    top_sig=OxmlElement("w:top"); top_sig.set(qn("w:val"),"single"); top_sig.set(qn("w:sz"),"6"); top_sig.set(qn("w:space"),"1"); top_sig.set(qn("w:color"),"334155"); pBdr_sig.append(top_sig); pPr_sig.append(pBdr_sig)
    r_sig=p_sig.add_run(f"\n{responsavel.nome.upper() if responsavel.nome else 'RESPONSÁVEL TÉCNICO'}"); r_sig.bold=True; r_sig.font.size=Pt(11); r_sig.font.color.rgb=PRETO; r_sig.font.name="Calibri"
    if responsavel.registro:
        p_reg=doc.add_paragraph(); p_reg.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r_reg=p_reg.add_run(f"Registro: {responsavel.registro}"); r_reg.font.size=Pt(9.5); r_reg.font.color.rgb=CINZA; r_reg.font.name="Calibri"
    p_tit=doc.add_paragraph(); p_tit.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r_tit=p_tit.add_run("Anotador de Responsabilidade Técnica"); r_tit.bold=True; r_tit.font.size=Pt(9); r_tit.font.color.rgb=CINZA_CLR; r_tit.font.name="Calibri"
    p_dt=doc.add_paragraph(); p_dt.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r_dt=p_dt.add_run(f"Relatório chancelado em: {hoje}"); r_dt.font.size=Pt(8.5); r_dt.font.color.rgb=CINZA_CLR; r_dt.font.name="Calibri"

    buf=BytesIO(); doc.save(buf); buf.seek(0)
    return buf.read()